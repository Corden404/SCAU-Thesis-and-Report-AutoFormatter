import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- 1. 手动定义缺失的枚举 (防止 ImportError) ---
class WD_OUTLINE_LEVEL:
    LEVEL_1 = 0
    LEVEL_2 = 1
    LEVEL_3 = 2
    LEVEL_4 = 3
    LEVEL_5 = 4
    LEVEL_6 = 5
    LEVEL_7 = 6
    LEVEL_8 = 7
    LEVEL_9 = 8
    BODY_TEXT = 9

# --- 2. 字体常量 (使用英文名以保证兼容性) ---
FONTS = {
    "zh_serif": "SimSun",           # 宋体
    "zh_sans":  "SimHei",           # 黑体
    "zh_kai":   "KaiTi",            # 楷体
    "en_serif": "Times New Roman"   # 西文
}

SIZES = {
    "二号": 22,
    "三号": 16,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9
}

def set_style_font(style, zh_font_name, en_font_name, size_pt, bold=False, italic=False):
    """
    设置字体：强制黑色、清除主题绑定、设置 Hint
    """
    # 1. 基本属性
    font = style.font
    font.name = en_font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(0, 0, 0) # 强制黑色

    # 2. XML 底层设置
    rPr = style.element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)

    # 获取或创建 rFonts
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    # 设置中西文字体
    rFonts.set(qn('w:ascii'), en_font_name)
    rFonts.set(qn('w:hAnsi'), en_font_name)
    rFonts.set(qn('w:cs'), en_font_name)
    rFonts.set(qn('w:eastAsia'), zh_font_name)

    # 3. 【关键】清除主题字体绑定 (防止 MS Gothic/等线 覆盖)
    if qn('w:asciiTheme') in rFonts.attrib:
        del rFonts.attrib[qn('w:asciiTheme')]
    if qn('w:eastAsiaTheme') in rFonts.attrib:
        del rFonts.attrib[qn('w:eastAsiaTheme')]
    if qn('w:cstheme') in rFonts.attrib:
        del rFonts.attrib[qn('w:cstheme')]

    # 4. 设置 Hint (优先东亚字体渲染)
    rFonts.set(qn('w:hint'), 'eastAsia')

def set_paragraph_format(style, align=None, line_spacing=1.5, 
                         first_indent_chars=0, hanging_indent_chars=0, 
                         left_indent_chars=0,
                         space_before=0, space_after=0, 
                         outline_level=None):
    pf = style.paragraph_format
    
    if align is not None:
        pf.alignment = align
    
    if line_spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif line_spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # 缩进计算
    font_size = style.font.size.pt if style.font.size else 12
    char_width = Pt(font_size)

    if first_indent_chars > 0:
        pf.first_line_indent = char_width * first_indent_chars
    else:
        pf.first_line_indent = Pt(0)

    if hanging_indent_chars > 0:
        pf.left_indent = char_width * hanging_indent_chars
        pf.first_line_indent = -1 * char_width * hanging_indent_chars
    elif left_indent_chars > 0:
        pf.left_indent = char_width * left_indent_chars
    else:
        pf.left_indent = Pt(0)

    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

    if outline_level is not None:
        try:
            pf.outline_level = outline_level
        except AttributeError:
            pass

def create_or_get_style(doc, style_name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles.add_style(style_name, style_type)

def main():
    doc = Document()
    print("正在生成 reference.docx ...")

    # ==========================================
    # 1. 基础与章节标题 (Built-in) - 【不加粗】
    # ==========================================
    
    # Normal: 宋体 / 小四
    style = doc.styles['Normal']
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小四'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, first_indent_chars=2)

    # Heading 1: 黑体 / 四号 / 不加粗
    style = doc.styles['Heading 1']
    set_style_font(style, FONTS['zh_sans'], FONTS['en_serif'], SIZES['四号'], bold=False)
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, first_indent_chars=0, outline_level=WD_OUTLINE_LEVEL.LEVEL_1)
    
    # Heading 2: 黑体 / 小四 / 不加粗
    style = doc.styles['Heading 2']
    set_style_font(style, FONTS['zh_sans'], FONTS['en_serif'], SIZES['小四'], bold=False)
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, first_indent_chars=0, outline_level=WD_OUTLINE_LEVEL.LEVEL_2)

    # Heading 3: 楷体 / 小四 / 不加粗
    style = doc.styles['Heading 3']
    set_style_font(style, FONTS['zh_kai'], FONTS['en_serif'], SIZES['小四'], bold=False)
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, first_indent_chars=0, outline_level=WD_OUTLINE_LEVEL.LEVEL_3)

    # Heading 4: 楷体 / 小四 / 不加粗
    style = doc.styles['Heading 4']
    set_style_font(style, FONTS['zh_kai'], FONTS['en_serif'], SIZES['小四'], bold=False)
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, first_indent_chars=0, outline_level=WD_OUTLINE_LEVEL.LEVEL_4)

    # ==========================================
    # 2. 目录样式 (TOC)
    # ==========================================
    for i in range(1, 4):
        try:
            style_name = f'TOC {i}'
            try:
                style = doc.styles[style_name]
            except KeyError:
                continue
            set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小四'])
            set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, first_indent_chars=0)
        except:
            pass

    # ==========================================
    # 3. 前置部分 (自定义样式)
    # ==========================================

    # SCAU_Abstract_Title: 黑体 / 四号 / 居中 / 不加粗
    style = create_or_get_style(doc, 'SCAU_Abstract_Title')
    set_style_font(style, FONTS['zh_sans'], FONTS['en_serif'], SIZES['四号'], bold=False)
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, outline_level=WD_OUTLINE_LEVEL.BODY_TEXT)

    # SCAU_Abstract_Body: 宋体 / 小四
    style = create_or_get_style(doc, 'SCAU_Abstract_Body')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小四'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, first_indent_chars=2)

    # SCAU_Keywords: 宋体 / 小四
    style = create_or_get_style(doc, 'SCAU_Keywords')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小四'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, first_indent_chars=0)

    # SCAU_English_Title: TNR / 四号 / 居中 / 加粗 (样式表特别指定了加粗)
    style = create_or_get_style(doc, 'SCAU_English_Title')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['四号'], bold=True)
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    # SCAU_Abstract_En: TNR / 小四
    style = create_or_get_style(doc, 'SCAU_Abstract_En')
    set_style_font(style, FONTS['en_serif'], FONTS['en_serif'], SIZES['小四'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, first_indent_chars=0)

    # ==========================================
    # 4. 后置部分 (自定义样式)
    # ==========================================

    # SCAU_Section_Centered (参考文献/致谢/附录标题): 黑体 / 四号 / 居中 / 不加粗
    style = create_or_get_style(doc, 'SCAU_Section_Centered')
    set_style_font(style, FONTS['zh_sans'], FONTS['en_serif'], SIZES['四号'], bold=False)
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, outline_level=WD_OUTLINE_LEVEL.LEVEL_1)

    # SCAU_References_Body: 宋体 / 小四 / 悬挂缩进
    style = create_or_get_style(doc, 'SCAU_References_Body')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小四'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, hanging_indent_chars=2)

    # SCAU_Ack_Body: 宋体 / 小四
    style = create_or_get_style(doc, 'SCAU_Ack_Body')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小四'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, first_indent_chars=2)

    # ==========================================
    # 5. 图表与组件
    # ==========================================

    # SCAU_Caption: 宋体 / 五号
    style = create_or_get_style(doc, 'SCAU_Caption')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['五号'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, first_indent_chars=0)

    # SCAU_Table_Content: 宋体 / 五号 / 单倍行距
    style = create_or_get_style(doc, 'SCAU_Table_Content')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['五号'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, first_indent_chars=0)

    # SCAU_Table_Note: 宋体 / 小五
    style = create_or_get_style(doc, 'SCAU_Table_Note')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小五'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, first_indent_chars=0)

    # SCAU_Footnote_Text: 宋体 / 小五 / 单倍行距
    style = create_or_get_style(doc, 'SCAU_Footnote_Text')
    set_style_font(style, FONTS['zh_serif'], FONTS['en_serif'], SIZES['小五'])
    set_paragraph_format(style, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0, first_indent_chars=0)

    # ==========================================
    # 6. 生成示例文本
    # ==========================================
    doc.add_paragraph("华南农业大学 SCAU Thesis Style Check", style="Heading 1")
    doc.add_paragraph("1.1 字体测试 (黑体不加粗)", style="Heading 2")
    doc.add_paragraph("1.1.1 楷体测试 (楷体不加粗)", style="Heading 3")
    
    p = doc.add_paragraph("正文应该是宋体。Body text should be SimSun.", style="Normal")
    
    p = doc.add_paragraph("参考文献 (黑体不加粗)", style="SCAU_Section_Centered")
    p = doc.add_paragraph("[1] Author. Title. Journal. 2024.", style="SCAU_References_Body")

    output_file = 'reference.docx'
    doc.save(output_file)
    print(f"成功生成: {output_file}")

if __name__ == "__main__":
    main()